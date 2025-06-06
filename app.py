import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
import tempfile
import os

st.set_page_config(page_title="Configuratore Attuatore", layout="centered")
st.title("⚙️ Configuratore Attuatore Elettromeccanico")

# Upload file
ciclo_file = st.file_uploader("📄 Carica ciclo posizione-tempo", type="xlsx")
viti_file = st.file_uploader("🔩 Database viti", type="xlsx")
motori_file = st.file_uploader("⚙️ Database motori", type="xlsx")
riduttori_file = st.file_uploader("🪛 Database riduttori", type="xlsx")

corsa_totale_input = st.number_input("📏 Corsa totale attuatore (mm)", min_value=1.0, value=150.0)
limite_acc = st.number_input("⏱️ Limite accelerazione (mm/s²)", value=5000.0)
limite_jerk = st.number_input("⚡ Limite jerk (mm/s³)", value=50000.0)

if all([ciclo_file, viti_file, motori_file, riduttori_file]):
    if st.button("▶️ Calcola"):
        df = pd.read_excel(ciclo_file)
        df.columns = [c.strip().lower() for c in df.columns]
        if not {'tempo', 'posizione', 'forza'}.issubset(df.columns):
            st.error("❌ Il file deve contenere le colonne: 'tempo', 'posizione', 'forza'")
            st.stop()

        df = df.sort_values("tempo")
        df["velocita"] = np.gradient(df["posizione"], df["tempo"])
        df["accelerazione"] = np.gradient(df["velocita"], df["tempo"])
        df["jerk"] = np.gradient(df["accelerazione"], df["tempo"])
        st.write("✅ Jerk calcolato")

        max_acc = df["accelerazione"].abs().max()
        max_jerk = df["jerk"].abs().max()

        st.subheader("📈 Analisi del ciclo")
        if max_acc > limite_acc:
            st.warning("⚠️ Accelerazione oltre il limite")
        if max_jerk > limite_jerk:
            st.warning("⚠️ Jerk oltre il limite")

        fig, axs = plt.subplots(4, 1, figsize=(8, 10), sharex=True)
        axs[0].plot(df["tempo"], df["posizione"]); axs[0].set_ylabel("Posizione [mm]")
        axs[1].plot(df["tempo"], df["velocita"]); axs[1].set_ylabel("Velocità [mm/s]")
        axs[2].plot(df["tempo"], df["accelerazione"]); axs[2].set_ylabel("Accelerazione [mm/s²]")
        axs[3].plot(df["tempo"], df["jerk"]); axs[3].set_ylabel("Jerk [mm/s³]"); axs[3].set_xlabel("Tempo [s]")
        st.pyplot(fig)

        st.subheader("📏 Verifica corsa")
        st.write("✅ Analisi corsa iniziata")
        corsa_ciclo = df["posizione"].max() - df["posizione"].min()
        st.write(f"Corsa effettiva: {corsa_ciclo:.1f} mm")
        if corsa_ciclo > corsa_totale_input:
            st.error("❌ Corsa richiesta superiore alla corsa disponibile")
            st.stop()

        st.subheader("🧮 Calcolo carico equivalente")
        st.write("✅ Carico equivalente calcolato")
        Feq = np.sqrt(np.mean(df["forza"]**2))
        st.write(f"Carico equivalente: {Feq:.0f} N")

        # Selezione vite
        st.subheader("🔩 Selezione vite")
        viti_df = pd.read_excel(viti_file)
        viti_valid = []
        for _, v in viti_df.iterrows():
            if Feq <= v['C'] and corsa_totale_input <= v['nocciolo'] * 25:
                viti_valid.append(v)
        if viti_valid:
            vite_valida = viti_valid[0]
            st.write("✅ Vite selezionata")
            st.success(f"Vite selezionata: {vite_valida['codice']}")
        else:
            st.error("❌ Nessuna vite compatibile trovata")
            st.stop()
        st.subheader("⚙️ Selezione riduttore")
        st.write("✅ File riduttori caricato")
        rid_df = pd.read_excel(riduttori_file)
        rid_df = pd.concat([pd.DataFrame([{"codice": "Diretta", "rapporto": 1.0, "rendimento": 1.0}]), rid_df], ignore_index=True)
        rid_sel = st.selectbox("Seleziona riduttore", rid_df["codice"])
        rid_r = rid_df[rid_df["codice"] == rid_sel]["rapporto"].values[0]
        rid_eta = rid_df[rid_df["codice"] == rid_sel]["rendimento"].values[0]

        # Coppia richiesta
        passo = vite_valida["passo"]
        rendimento_vite = vite_valida["rendimento"]
        torque_asse = Feq * passo / (2 * np.pi * rendimento_vite)
        torque_motore = torque_asse / (rid_r * rid_eta)
        rpm_asse = df["velocita"].max() / passo * 60 / 1000
        rpm_motore = rpm_asse * rid_r

        st.write(f"Coppia richiesta al motore: {torque_motore:.2f} Nm")
        st.write(f"Velocità motore richiesta: {rpm_motore:.0f} rpm")

        # Selezione motore
        st.subheader("🔌 Selezione motore")
        st.write("✅ File motori caricato")
        motori_df = pd.read_excel(motori_file)

        # Caricamento curve caratteristiche dei motori
        curve_files = st.file_uploader("Carica le curve motore (più file .xlsx)", type="xlsx", accept_multiple_files=True)
        curve_motori = {}
        if curve_files:
            for file in curve_files:
                try:
                    df_curve = pd.read_excel(file)
                    codice = file.name.replace('.xlsx','')
                    curve_motori[codice] = df_curve
                    st.success(f"✅ Caricata curva: {codice}")
                except Exception as e:
                    st.error(f"Errore nel file {file.name}: {e}")
        motori_validi = motori_df[
            (motori_df["coppia_massima"] >= torque_motore) &
            (motori_df["velocita_nominale"] >= rpm_motore)
        ]
        if not motori_validi.empty:
            st.success(f"Motore selezionato: {motori_validi.iloc[0]['codice']}")

        # Verifica curva motore e generazione grafico
        if codice_motore in curve_motori:
            curva = curve_motori[codice_motore]
            fig, ax = plt.subplots()
            ax.plot(curva['velocità'], curva['coppia_nominale'], label='Coppia Nominale')
            ax.plot(curva['velocità'], curva['coppia_massima'], label='Coppia Massima', linestyle='--')
            ax.set_title(f"Curva motore {codice_motore}")
            ax.set_xlabel('Velocità [rpm]')
            ax.set_ylabel('Coppia [Nm]')
            ax.legend()
            st.pyplot(fig)
            fig.savefig(f"curva_{codice_motore}.png")  # salva per report
        else:
            st.warning(f"Nessuna curva trovata per il motore {codice_motore}")
        else:
            st.warning("❌ Nessun motore compatibile trovato")

        # Vita utile
        st.subheader("🕒 Calcolo vita utile")
        C = vite_valida["C"]
        C0 = vite_valida["C0"]
        st.write("✅ Vita utile calcolata")
        vita_cicli = (C / Feq)**3 * 1e6
        vita_anni = vita_cicli / (3600 * 16 * 250)
        st.write(f"Cicli raggiungibili: {vita_cicli:,.0f}")
        st.write(f"Durata stimata: {vita_anni:.1f} anni (16h/giorno, 250gg/anno)")

        # Report DOCX
        if st.button("📤 Esporta report"):
            doc = Document()
            doc.add_heading("Report di Dimensionamento Attuatore", 0)
            doc.add_heading("📈 Ciclo", level=1)
            doc.add_paragraph(f"Max accelerazione: {max_acc:.1f} mm/s²")
            doc.add_paragraph(f"Max jerk: {max_jerk:.1f} mm/s³")
            doc.add_heading("📏 Verifiche", level=1)
            doc.add_paragraph(f"Corsa effettiva: {corsa_ciclo:.1f} mm (max: {corsa_totale_input} mm)")
            doc.add_paragraph(f"Carico equivalente: {Feq:.0f} N")
            doc.add_heading("⚙️ Componenti", level=1)
            doc.add_paragraph(f"Vite: {vite_valida['codice']}")
            doc.add_paragraph(f"Riduttore: {rid_sel}")
            if not motori_validi.empty:
                doc.add_paragraph(f"Motore: {motori_validi.iloc[0]['codice']}")
            doc.add_paragraph(f"RPM motore: {rpm_motore:.0f}, Coppia: {torque_motore:.2f} Nm")
            doc.add_heading("🕒 Vita utile", level=1)
            doc.add_paragraph(f"Cicli: {vita_cicli:,.0f}, Anni: {vita_anni:.1f}")
            temp_path = os.path.join(tempfile.gettempdir(), "report_dimensionamento.docx")
            doc.save(temp_path)
        try:
            doc.add_picture(f"curva_{codice_motore}.png", width=Inches(5.5))
        except Exception as e:
            st.warning("Curva non inserita nel report: " + str(e))
            with open(temp_path, "rb") as file:
                st.download_button("📄 Scarica Report DOCX", file, file_name="report_dimensionamento.docx")
